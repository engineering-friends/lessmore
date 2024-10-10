import os.path

from datetime import datetime

from docx import Document
from docx2pdf import convert
from lessmore.utils.file_primitives.ensure_path import ensure_path
from lessmore.utils.loguru_utils.setup_json_loguru import setup_json_loguru
from lessmore.utils.system.open_in_os import open_in_os
from loguru import logger
from markdown_it.rules_block import paragraph


def generate_docx_and_pdf(
    template_filename: str,
    replacements: dict,
    output_docx: str = "output.docx",
):
    # - Load the DOCX template

    doc = Document(template_filename)

    # - Replace placeholders with actual values

    # -- Collect runs from paragraphs and tables

    runs = []

    runs += sum([paragraph.runs for paragraph in doc.paragraphs], [])
    runs += sum(
        [
            paragraph.runs
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
            for paragraph in cell.paragraphs
        ],
        [],
    )

    # -- Fix placeholders

    for run in runs:
        try:
            run.text = run.text.format(**replacements)
        except:
            logger.exception("Failed to replace placeholders", text=run.text)
            raise

    # - Save the document

    doc.save(ensure_path(output_docx))

    # - Convert to pdf

    convert(output_docx, output_docx.replace(".docx", ".pdf"))


if __name__ == "__main__":
    # - Setup loguru

    setup_json_loguru()

    # - Init output path

    output_path = "data/2024-09 act.docx"

    # - Set current date

    now = datetime.now()

    # - Generate docx and pdf

    generate_docx_and_pdf(
        template_filename=os.path.abspath("generate_act_template.docx"),
        output_docx=output_path,
        replacements={
            # - Unchanged
            "ID": "105550154",
            "SERVICE_AGREEMENT": "Service Agreement No1-09/24 from 02.09.2024",
            # - Auto
            "TODAY_MM_YY": now.strftime("%m/%y"),
            "TODAY_YYYY_MM_DD": now.strftime("%Y-%m-%d"),
            # - Changed
            "N": 20,
            "HOURS": int(7500 / 50),
            "AMOUNT": 7500,
            "AMOUNT_WORDS": "seven hundred and fifty",
            "PAID_MONTH_YYYY_MM": "2024-09",
        },
    )

    # - Open pdf

    open_in_os(output_path.replace(".docx", ".pdf"))
